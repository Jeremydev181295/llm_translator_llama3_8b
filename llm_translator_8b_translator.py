from llama_cpp import Llama
from spire.doc import *
from spire.doc.common import *
import time
from module import handle_docx
from module import handle_footnote
import docx
from docx2python import docx2python

llm = Llama(
    model_path="models/Meta-Llama-3-8B-Instruct-Q8_0.gguf",
    n_ctx=4096,      # Max tokens for in + out
    n_threads=4,     # CPU cores used
    n_gpu_layers=-1,  # Load all layers into VRAM of the GPU
    
)

source_file_path = 'test.docx'

trans_lang = "English"
source_lang = "French"
style = "written"

def trans_with_ai(prompt, max_tokens=2048):
    """
    Function to send a prompt to the AI and return its response.
    """
    # This function sends the prompt to your AI model and fetches the response
    response = llm(prompt, max_tokens=max_tokens, temperature = 0.001, stop=["Q:", "\n"], echo=False)
    return response

def build_trans_prompt(source_lang, trans_lang, string):   
    prompt = f"Translate '{string.strip()}' from {source_lang} into {trans_lang}. I need only translation. Do not translate numbers, symbols and abbreviations. Keep original style. Keep dash and colon. "
    return prompt

def build_footnote_content_trans_prompt(source_lang, trans_lang, string):  
    prompt = f"Acts as a smart translator. Translate {source_lang} sentences into {trans_lang} sentences in written style. Do not remove heading word. Do not add any characters. If sentence includes '----footnotes----' then translate it. I need only translation sentence. '{string.strip()}'"
    return prompt

def build_footnote_reference_trans_prompt(source_lang, trans_lang, data):    
    prompts = [f"Acts as a smart translator. Translate {source_lang} sentences into {trans_lang} sentences in written style. Do not remove heading word. Do not add any characters. I need only translation sentence. '{string.strip()}'" for string in data if string.strip()]
    return prompts


def main():    
    
    document = Document()
    document.LoadFromFile(source_file_path)

    read_doc = docx2python(source_file_path)
    
    start_time = time.time()

    footnote_para_indexes = handle_footnote.find_paragraphs_for_footnote(file_path=source_file_path)
    unique_footnote_para_indexes = list(set(footnote_para_indexes))
    if footnote_para_indexes != []:
        print("footnote paragraphs indexes:", footnote_para_indexes)
        footnote_para_strings = handle_docx.extract_footnote_para_strings(doc=read_doc, index_list=unique_footnote_para_indexes)
    ############################################################################################################
    # translate content, header and footer 
    for index in range(len(document.Sections)):
        section = document.Sections[index]

        # translate content
        for i in range(len(section.Paragraphs)):
            if i in unique_footnote_para_indexes:
                string = footnote_para_strings[unique_footnote_para_indexes.index(i)]
                if not string.strip() =='':
                        prompt = build_footnote_content_trans_prompt(source_lang=source_lang, trans_lang=trans_lang, string=string)
                        prompt = f"Q: {prompt} A: "
                        ai_response = trans_with_ai(prompt)           
                        print(ai_response)
                        section.Paragraphs[i].Text = ai_response['choices'][0]['text'].strip()
            else:
                string = section.Paragraphs[i].Text
                if not string.strip() =='':
                    if not 'https://' in string:
                        prompt = build_trans_prompt(source_lang=source_lang, trans_lang=trans_lang, string=string)
                        prompt = f"Q: {prompt} A: "
                        ai_response = trans_with_ai(prompt)           
                        print(ai_response)
                        section.Paragraphs[i].Text = ai_response['choices'][0]['text'].strip()
            


        # translate header
        for i in range(len(section.HeadersFooters.Header.Paragraphs)):
            string = section.HeadersFooters.Header.Paragraphs[i].Text
            if not string.strip() =='':
                prompt = build_trans_prompt(source_lang=source_lang, trans_lang=trans_lang, string=string)
                prompt = f"Q: {prompt} A: "
                ai_response = trans_with_ai(prompt)           
                print(ai_response)
                section.HeadersFooters.Header.Paragraphs[i].Text = ai_response['choices'][0]['text'].strip()

        # translate footer
        for i in range(len(section.HeadersFooters.Footer.Paragraphs)):
            string = section.HeadersFooters.Footer.Paragraphs[i].Text
            if not string.strip() =='':
                prompt = build_trans_prompt(source_lang=source_lang, trans_lang=trans_lang, string=string)
                prompt = f"Q: {prompt} A: "
                ai_response = trans_with_ai(prompt)           
                print(ai_response)
                section.HeadersFooters.Footer.Paragraphs[i].Text = ai_response['choices'][0]['text'].strip()

    content_translated_file_path = 'content_header_footer_translated_result.docx'
    document.SaveToFile(content_translated_file_path, FileFormat.Docx)

    ###########################################################################################################
    # remove unnecessary first paragraph
    revise_doc = docx.Document(content_translated_file_path)
    if revise_doc.paragraphs[0].text == "Evaluation Warning: The document was created with Spire.Doc for Python.":
            handle_docx.delete_paragraph(revise_doc.paragraphs[0])
    revise_doc.save(content_translated_file_path)

    ###########################################################################################################
    # handling footnotes
    if footnote_para_indexes != []:
        # build footnote prompts   
        footnote_data = handle_docx.extract_footnote(read_doc)
        if footnote_data:
            footnote_prompts = build_footnote_reference_trans_prompt(source_lang=source_lang, trans_lang=trans_lang, data=footnote_data)
        else:
            footnote_prompts = ""

        # translate footnote of file
        result_footnote_data = []
        if footnote_prompts:
            for prompt in footnote_prompts:
                # prompt = system_prompt_footnote + prompt
                prompt = f"Q: {prompt} A: "
                ai_response = trans_with_ai(prompt=prompt)
                print(ai_response)
                result_footnote_data.append(ai_response['choices'][0]['text'].strip())

        # add foottnotes
        refer_string = []
        for i, index in enumerate(footnote_para_indexes):
            doc = docx2python(content_translated_file_path)
            refer_string.append(handle_docx.extract_footnote_pos_string(doc, '----footnote'+str(i+1)+'----'))
            print(refer_string)

            # remove footnote strings
            remove_footnote_string_doc = docx.Document(content_translated_file_path)
            handle_docx.remove_string_from_paragraph(remove_footnote_string_doc, '----footnote'+str(i+1)+'----')     
            remove_footnote_string_doc.save(content_translated_file_path)

        for i, index in enumerate(footnote_para_indexes):
            
            handle_footnote.add_footnote(file_path=content_translated_file_path, para_index=index, refer_string=refer_string[i], footnote_text=result_footnote_data[i])
            
            # remove unnecessary string created from using spire.doc package    
            revise_doc = docx.Document(content_translated_file_path)
            if revise_doc.paragraphs[0].text == "Evaluation Warning: The document was created with Spire.Doc for Python.":
                    handle_docx.delete_paragraph(revise_doc.paragraphs[0])
        
        revise_doc.save(content_translated_file_path)
    
    revise_doc = docx.Document(content_translated_file_path)
    handle_docx.remove_string_from_paragraph(revise_doc, '----footnotes----')     
    revise_doc.save(content_translated_file_path)
    
    ###########################################################################################################
    # translate table
    output_path = 'final_result.docx'
    # Load the DOCX file
    doc = docx.Document(content_translated_file_path)
    
    # Iterate through all tables in the document
    for table in doc.tables:
        # Iterate through each row in the table
        for row in table.rows:
            # Iterate through each cell in the row
            for cell in row.cells:
                if cell.text:
                    prompt = build_trans_prompt(source_lang=source_lang, trans_lang=trans_lang, string=cell.text)
                    prompt = f"Q: {prompt} A: "
                    cell.text = trans_with_ai(prompt)['choices'][0]['text'].strip()
                    print(cell.text)
    # Save the modified document
    doc.save(output_path)

    ###########################################################################################################

    end_time = time.time()
    elapsed_time = end_time - start_time
    print(elapsed_time)
   
# To start the chat, call the main_chat function
main()